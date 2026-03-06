"""
Extract text from .docx files.
"""
import io

try:
    from docx import Document
except ImportError:
    Document = None


def extract_text_from_docx(file_stream) -> str:
    """
    Read DOCX from file-like object. Returns full document text.
    """
    if Document is None:
        raise RuntimeError("python-docx required for DOCX extraction. pip install python-docx")

    doc = Document(file_stream)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)


def extract_text_from_docx_bytes(data: bytes) -> str:
    return extract_text_from_docx(io.BytesIO(data))
